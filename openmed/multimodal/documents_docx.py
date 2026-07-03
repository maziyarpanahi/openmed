"""DOCX text extraction and run-level write-back for multimodal redaction.

The DOCX ingester imports ``python-docx`` lazily so importing
``openmed.multimodal`` stays lightweight. Extraction emits normalized text plus
one source span per non-empty Word run, including paragraphs in headers,
footers, and table cells. Redaction helpers project detected text spans back to
the covered runs and can write a clean DOCX while preserving document structure.
Tracked changes and comment redaction remain out of scope for this ingester.
"""

from __future__ import annotations

import hashlib
import importlib
from collections import defaultdict
from collections.abc import Iterable, Mapping, Sequence
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from .base import ExtractedDocument, SourceSpan, register_handler
from .exceptions import MissingDependencyError

_DOCX_HINT = 'Install with: pip install "openmed[multimodal]".'
_BLOCK_SEPARATOR = "\n"


@dataclass(frozen=True)
class DocxRunRange:
    """A local run range covered by a detected normalized-text span."""

    document_run_index: int
    run_start: int
    run_end: int
    metadata: Mapping[str, Any] = field(default_factory=dict)

    def to_dict(self) -> dict[str, Any]:
        """Return a PHI-safe metadata representation."""
        return {
            "document_run_index": self.document_run_index,
            "run_start": self.run_start,
            "run_end": self.run_end,
            "metadata": dict(self.metadata),
        }


@dataclass(frozen=True)
class DocxRedaction:
    """A detected PHI span projected to one or more DOCX runs."""

    start: int
    end: int
    replacement: str
    label: str | None = None
    confidence: float | None = None
    run_ranges: tuple[DocxRunRange, ...] = ()
    metadata: Mapping[str, Any] = field(default_factory=dict)

    def to_dict(self) -> dict[str, Any]:
        """Return a PHI-safe metadata representation."""
        payload: dict[str, Any] = {
            "start": self.start,
            "end": self.end,
            "replacement": self.replacement,
            "run_ranges": [run_range.to_dict() for run_range in self.run_ranges],
        }
        if self.label is not None:
            payload["label"] = self.label
        if self.confidence is not None:
            payload["confidence"] = self.confidence
        if self.metadata:
            payload["metadata"] = dict(self.metadata)
        return payload


@dataclass(frozen=True)
class _ParagraphContext:
    paragraph: Any
    part: str
    block_type: str
    paragraph_index: int
    section_index: int | None = None
    table_index: int | None = None
    row_index: int | None = None
    cell_index: int | None = None
    container_path: str | None = None


@dataclass
class _ExtractionState:
    parts: list[str] = field(default_factory=list)
    spans: list[SourceSpan] = field(default_factory=list)
    cursor: int = 0
    block_index: int = 0
    document_run_index: int = 0
    text_run_count: int = 0
    paragraph_count: int = 0


@dataclass(frozen=True)
class _RunEdit:
    start: int
    end: int
    replacement: str


def _import_docx() -> Any:
    try:
        return importlib.import_module("docx")
    except ImportError as exc:  # pragma: no cover - exercised without extra.
        raise MissingDependencyError(
            dependency="python-docx", instruction=_DOCX_HINT
        ) from exc


def _import_docx_types() -> tuple[Any, Any, Any, Any]:
    try:
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError as exc:  # pragma: no cover - exercised without extra.
        raise MissingDependencyError(
            dependency="python-docx", instruction=_DOCX_HINT
        ) from exc
    return CT_P, CT_Tbl, Paragraph, Table


def extract_docx(path: str | Path) -> ExtractedDocument:
    """Extract normalized DOCX text plus char-offset source spans.

    Paragraphs are separated with newlines. Each non-empty Word run becomes a
    :class:`SourceSpan` whose metadata includes the document-wide run index and
    paragraph/table/header/footer location needed for write-back.
    """
    docx = _import_docx()
    word_document = docx.Document(path)
    state = _ExtractionState()

    for context in _iter_docx_paragraphs(word_document):
        _append_paragraph(context, state)

    return ExtractedDocument(
        text="".join(state.parts),
        spans=tuple(state.spans),
        metadata={
            "format": "docx",
            "paragraph_count": state.paragraph_count,
            "text_run_count": state.text_run_count,
            "document_run_count": state.document_run_index,
        },
    )


def map_text_spans_to_docx_runs(
    document: ExtractedDocument,
    spans: Iterable[Any] | Mapping[str, Any] | Any,
    *,
    replacement: str | None = None,
) -> tuple[DocxRedaction, ...]:
    """Project detected normalized-text spans to DOCX run ranges.

    ``spans`` accepts the same common detector shapes used by the PDF handler:
    mappings or objects with ``start``/``end``, ``(start, end)`` tuples, or a
    detector result exposing ``entities``, ``pii_entities`` or ``spans``.
    """
    redactions: list[DocxRedaction] = []
    for raw_entity in _iter_entity_inputs(spans):
        entity = _coerce_entity(raw_entity, default_replacement=replacement)
        if entity is None:
            continue
        start, end, label, confidence, entity_replacement = entity
        start = max(0, start)
        end = min(len(document.text), end)
        if end <= start:
            continue

        covered = tuple(
            source
            for source in document.spans
            if source.metadata.get("format") == "docx"
            and source.end > start
            and source.start < end
        )
        if not covered:
            continue

        run_ranges = tuple(
            _source_span_to_run_range(source, start=start, end=end)
            for source in covered
        )
        redactions.append(
            DocxRedaction(
                start=start,
                end=end,
                label=label,
                confidence=confidence,
                replacement=entity_replacement,
                run_ranges=run_ranges,
                metadata={
                    "text_sha256": hashlib.sha256(
                        document.text[start:end].encode("utf-8")
                    ).hexdigest(),
                    "source_span_count": len(run_ranges),
                },
            )
        )
    return tuple(redactions)


def write_redacted_docx(
    source_path: str | Path,
    output_path: str | Path,
    spans: Iterable[Any] | Mapping[str, Any] | Any,
    *,
    replacement: str | None = None,
) -> Path:
    """Write a redacted DOCX by applying normalized-text spans to source runs."""
    document = extract_docx(source_path)
    redactions = map_text_spans_to_docx_runs(document, spans, replacement=replacement)
    _write_docx_redactions(source_path, output_path, redactions)
    return Path(output_path)


def _append_paragraph(context: _ParagraphContext, state: _ExtractionState) -> None:
    paragraph_had_text = False
    block_index = state.block_index

    for run_index, run in enumerate(context.paragraph.runs):
        document_run_index = state.document_run_index
        state.document_run_index += 1

        text = str(run.text)
        if not text:
            continue

        if not paragraph_had_text:
            if state.parts:
                state.parts.append(_BLOCK_SEPARATOR)
                state.cursor += len(_BLOCK_SEPARATOR)
            paragraph_had_text = True

        start = state.cursor
        state.parts.append(text)
        state.cursor += len(text)
        state.spans.append(
            SourceSpan(
                start=start,
                end=state.cursor,
                metadata=_source_metadata(
                    context,
                    block_index=block_index,
                    run_index=run_index,
                    document_run_index=document_run_index,
                ),
            )
        )
        state.text_run_count += 1

    if paragraph_had_text:
        state.block_index += 1
        state.paragraph_count += 1


def _source_metadata(
    context: _ParagraphContext,
    *,
    block_index: int,
    run_index: int,
    document_run_index: int,
) -> dict[str, Any]:
    metadata: dict[str, Any] = {
        "format": "docx",
        "block_type": context.block_type,
        "part": context.part,
        "block_index": block_index,
        "paragraph_index": context.paragraph_index,
        "run_index": run_index,
        "document_run_index": document_run_index,
    }
    optional_values = {
        "section_index": context.section_index,
        "table_index": context.table_index,
        "row_index": context.row_index,
        "cell_index": context.cell_index,
        "container_path": context.container_path,
    }
    metadata.update(
        {key: value for key, value in optional_values.items() if value is not None}
    )
    return metadata


def _iter_docx_paragraphs(word_document: Any) -> Iterable[_ParagraphContext]:
    paragraph_index = 0
    table_index = 0
    seen_header_footer: set[int] = set()

    def next_paragraph(
        paragraph: Any,
        *,
        part: str,
        block_type: str,
        section_index: int | None = None,
        table: int | None = None,
        row: int | None = None,
        cell: int | None = None,
        path: str | None = None,
    ) -> _ParagraphContext:
        nonlocal paragraph_index
        context = _ParagraphContext(
            paragraph=paragraph,
            part=part,
            block_type=block_type,
            section_index=section_index,
            paragraph_index=paragraph_index,
            table_index=table,
            row_index=row,
            cell_index=cell,
            container_path=path,
        )
        paragraph_index += 1
        return context

    def walk_container(
        container: Any,
        *,
        part: str,
        section_index: int | None,
        path: str,
        in_table_cell: bool = False,
        table: int | None = None,
        row: int | None = None,
        cell: int | None = None,
    ) -> Iterable[_ParagraphContext]:
        nonlocal table_index
        _, _, paragraph_type, table_type = _import_docx_types()
        for block in _iter_block_items(container):
            if isinstance(block, paragraph_type):
                block_type = (
                    "table_cell"
                    if in_table_cell
                    else "paragraph"
                    if part == "body"
                    else part
                )
                yield next_paragraph(
                    block,
                    part=part,
                    block_type=block_type,
                    section_index=section_index,
                    table=table,
                    row=row,
                    cell=cell,
                    path=path,
                )
            elif isinstance(block, table_type):
                current_table = table_index
                table_index += 1
                for row_index, table_row in enumerate(block.rows):
                    for cell_index, table_cell in enumerate(table_row.cells):
                        cell_path = (
                            f"{path}.table{current_table}.r{row_index}.c{cell_index}"
                        )
                        yield from walk_container(
                            table_cell,
                            part=part,
                            section_index=section_index,
                            path=cell_path,
                            in_table_cell=True,
                            table=current_table,
                            row=row_index,
                            cell=cell_index,
                        )

    for section_index, section in enumerate(word_document.sections):
        container = section.header
        container_id = id(getattr(container, "_element", container))
        if container_id not in seen_header_footer:
            seen_header_footer.add(container_id)
            yield from walk_container(
                container,
                part="header",
                section_index=section_index,
                path=f"header{section_index}",
            )

    yield from walk_container(
        word_document,
        part="body",
        section_index=None,
        path="body",
    )

    for section_index, section in enumerate(word_document.sections):
        container = section.footer
        container_id = id(getattr(container, "_element", container))
        if container_id not in seen_header_footer:
            seen_header_footer.add(container_id)
            yield from walk_container(
                container,
                part="footer",
                section_index=section_index,
                path=f"footer{section_index}",
            )


def _iter_block_items(parent: Any) -> Iterable[Any]:
    ct_p, ct_tbl, paragraph_type, table_type = _import_docx_types()
    parent_element = _container_element(parent)
    for child in parent_element.iterchildren():
        if isinstance(child, ct_p):
            yield paragraph_type(child, parent)
        elif isinstance(child, ct_tbl):
            yield table_type(child, parent)


def _container_element(parent: Any) -> Any:
    element = getattr(parent, "element", None)
    if element is not None and hasattr(element, "body"):
        return element.body
    cell_element = getattr(parent, "_tc", None)
    if cell_element is not None:
        return cell_element
    private_element = getattr(parent, "_element", None)
    if private_element is not None:
        return private_element
    if element is not None:
        return element
    raise TypeError(f"Unsupported DOCX container {type(parent)!r}")


def _source_span_to_run_range(
    source: SourceSpan,
    *,
    start: int,
    end: int,
) -> DocxRunRange:
    metadata = dict(source.metadata)
    return DocxRunRange(
        document_run_index=int(metadata["document_run_index"]),
        run_start=max(start, source.start) - source.start,
        run_end=min(end, source.end) - source.start,
        metadata={
            key: value
            for key, value in metadata.items()
            if key
            in {
                "block_type",
                "part",
                "block_index",
                "paragraph_index",
                "run_index",
                "section_index",
                "table_index",
                "row_index",
                "cell_index",
                "container_path",
            }
        },
    )


def _write_docx_redactions(
    source_path: str | Path,
    output_path: str | Path,
    redactions: Sequence[DocxRedaction],
) -> None:
    docx = _import_docx()
    word_document = docx.Document(source_path)
    runs = _collect_runs(word_document)
    edits_by_run: dict[int, list[_RunEdit]] = defaultdict(list)

    for redaction in _non_overlapping_redactions(redactions):
        replacement_inserted = False
        for run_range in redaction.run_ranges:
            replacement = "" if replacement_inserted else redaction.replacement
            replacement_inserted = True
            edits_by_run[run_range.document_run_index].append(
                _RunEdit(
                    start=run_range.run_start,
                    end=run_range.run_end,
                    replacement=replacement,
                )
            )

    for document_run_index, edits in edits_by_run.items():
        run = runs.get(document_run_index)
        if run is None:
            continue
        text = str(run.text)
        for edit in sorted(edits, key=lambda item: item.start, reverse=True):
            text = text[: edit.start] + edit.replacement + text[edit.end :]
        run.text = text

    word_document.save(output_path)


def _collect_runs(word_document: Any) -> dict[int, Any]:
    runs: dict[int, Any] = {}
    document_run_index = 0
    for context in _iter_docx_paragraphs(word_document):
        for run in context.paragraph.runs:
            runs[document_run_index] = run
            document_run_index += 1
    return runs


def _non_overlapping_redactions(
    redactions: Sequence[DocxRedaction],
) -> tuple[DocxRedaction, ...]:
    selected: list[DocxRedaction] = []
    cursor = -1
    for redaction in sorted(
        redactions, key=lambda item: (item.start, -(item.end - item.start), item.end)
    ):
        if redaction.start < cursor:
            continue
        selected.append(redaction)
        cursor = redaction.end
    return tuple(selected)


def _detect_entities(document: ExtractedDocument, models: Any, lang: str | None) -> Any:
    detector = _resolve_detector(models)
    if detector is None:
        return ()
    try:
        return detector(document.text, lang=lang)
    except TypeError:
        return detector(document.text)


def _resolve_detector(models: Any) -> Any:
    if models is None:
        return None
    if callable(models):
        return models
    if isinstance(models, Mapping):
        for key in ("detector", "extract_pii", "analyze_text", "predict_entities"):
            candidate = models.get(key)
            if callable(candidate):
                return candidate
        return None
    for name in (
        "detect",
        "extract_pii",
        "analyze_text",
        "predict_entities",
        "predict",
    ):
        candidate = getattr(models, name, None)
        if callable(candidate):
            return candidate
    return None


def _iter_entity_inputs(spans: Any) -> tuple[Any, ...]:
    if spans is None:
        return ()
    entities = getattr(spans, "entities", None)
    if entities is not None:
        return tuple(entities)
    pii_entities = getattr(spans, "pii_entities", None)
    if pii_entities is not None:
        return tuple(pii_entities)
    if isinstance(spans, Mapping):
        for key in ("entities", "pii_entities", "spans"):
            entities = spans.get(key)
            if entities is not None:
                return tuple(entities)
        if "start" in spans and "end" in spans:
            return (spans,)
    if _looks_like_sequence_entity(spans):
        return (spans,)
    if isinstance(spans, Iterable) and not isinstance(spans, (str, bytes, bytearray)):
        return tuple(spans)
    return ()


def _coerce_entity(
    span: Any,
    *,
    default_replacement: str | None,
) -> tuple[int, int, str | None, float | None, str] | None:
    if _looks_like_sequence_entity(span):
        label = _coerce_optional_str(span[2]) if len(span) >= 3 else None
        return (
            int(span[0]),
            int(span[1]),
            label,
            None,
            default_replacement or _mask_for_label(label),
        )

    if isinstance(span, Mapping):
        start = span.get("start")
        end = span.get("end")
        label = span.get("label", span.get("entity_type", span.get("entity_group")))
        confidence = span.get("confidence", span.get("score"))
        replacement = span.get("replacement", span.get("redacted_text"))
    else:
        start = getattr(span, "start", None)
        end = getattr(span, "end", None)
        label = getattr(
            span,
            "label",
            getattr(span, "entity_type", getattr(span, "entity_group", None)),
        )
        confidence = getattr(span, "confidence", getattr(span, "score", None))
        replacement = getattr(span, "replacement", getattr(span, "redacted_text", None))

    if start is None or end is None:
        return None
    label_text = _coerce_optional_str(label)
    replacement_text = _coerce_optional_str(replacement)
    return (
        int(start),
        int(end),
        label_text,
        _coerce_confidence(confidence),
        replacement_text or default_replacement or _mask_for_label(label_text),
    )


def _looks_like_sequence_entity(value: Any) -> bool:
    if isinstance(value, Sequence) and not isinstance(value, (str, bytes, bytearray)):
        if len(value) >= 2:
            try:
                int(value[0])
                int(value[1])
            except (TypeError, ValueError):
                return False
            return True
    return False


def _coerce_optional_str(value: Any) -> str | None:
    if value is None:
        return None
    text = str(value)
    return text or None


def _coerce_confidence(value: Any) -> float | None:
    if value is None:
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _mask_for_label(label: str | None) -> str:
    safe_label = "".join(
        character if character.isalnum() else "_"
        for character in (label or "PHI").upper()
    ).strip("_")
    return f"[{safe_label or 'PHI'}]"


def _policy_value(policy: Any, *names: str) -> Any:
    if policy is None:
        return None
    if isinstance(policy, Mapping):
        for name in names:
            if name in policy:
                return policy[name]
        return None
    for name in names:
        value = getattr(policy, name, None)
        if value is not None:
            return value
    return None


def _docx_handler(
    path: str | Path,
    *,
    policy: Any = None,
    models: Any = None,
    lang: str | None = None,
) -> ExtractedDocument:
    document = extract_docx(path)
    entities = _iter_entity_inputs(_detect_entities(document, models, lang))
    replacement = _coerce_optional_str(_policy_value(policy, "replacement"))
    redactions = map_text_spans_to_docx_runs(
        document, entities, replacement=replacement
    )
    if not redactions:
        return document

    output_path = _policy_value(
        policy, "output_path", "redacted_path", "destination_path"
    )
    if output_path is not None:
        _write_docx_redactions(path, output_path, redactions)

    metadata = dict(document.metadata)
    metadata.update(
        {
            "detected_span_count": len(entities),
            "docx_redactions": [redaction.to_dict() for redaction in redactions],
        }
    )
    if output_path is not None:
        metadata["redacted_docx_path"] = str(output_path)
    if policy is not None:
        metadata["policy"] = policy

    return ExtractedDocument(
        text=document.text,
        spans=document.spans,
        metadata=metadata,
    )


register_handler(".docx", _docx_handler)


__all__ = [
    "DocxRedaction",
    "DocxRunRange",
    "extract_docx",
    "map_text_spans_to_docx_runs",
    "write_redacted_docx",
]
